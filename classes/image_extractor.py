import os
import shutil
import uuid
import zipfile
import tempfile
import time
import datetime
from pathlib import Path
from typing import List, Dict, Any
from PIL import Image
import fitz  # PyMuPDF for PDF image extraction
from docx import Document
from docx.image.exceptions import UnrecognizedImageError
import xml.etree.ElementTree as ET
from classes.models import ImageInfo
from classes import config

class ImageExtractor:
    """Extract images from various document types and save them to accessible folders"""

    def __init__(self, base_url: str = None, images_dir: str = None):
        # Use config.IMAGE_BASE_URL if no specific base_url is provided
        self.base_url = (base_url or config.IMAGE_BASE_URL).rstrip('/')
        # Use config.IMAGES_DIR if no specific directory is provided
        self.images_dir = Path(images_dir or config.IMAGES_DIR)
        # Don't create directory immediately - do it lazily when first needed

    def _ensure_images_dir_exists(self):
        """Ensure the images directory exists, create if necessary"""
        try:
            self.images_dir.mkdir(parents=True, exist_ok=True, mode=0o755)
        except PermissionError:
            # If we can't create the directory, try to use a temporary location
            # This should align with the fallback used in main.py
            import tempfile
            temp_dir = Path(tempfile.gettempdir()) / "markitdown_static" / "images"
            temp_dir.mkdir(parents=True, exist_ok=True, mode=0o755)

            # Update the base_url to reflect the temporary directory
            old_images_dir = self.images_dir
            self.images_dir = temp_dir
            print(f"Warning: Could not create {old_images_dir}, using temporary directory: {temp_dir}")
            print(f"Images will be served from temporary location")

    def extract_images_from_file(self, file_path: str, document_name: str) -> List[ImageInfo]:
        """Extract images from a file based on its extension"""
        print(f"DEBUG: ImageExtractor.extract_images_from_file called with:")
        print(f"DEBUG: - file_path: {file_path}")
        print(f"DEBUG: - document_name: {document_name}")

        # Ensure directory exists before extraction
        self._ensure_images_dir_exists()

        file_path = Path(file_path)
        document_folder = self._create_document_folder(document_name)

        print(f"DEBUG: Created document folder: {document_folder}")

        if file_path.suffix.lower() == '.pdf':
            return self._extract_from_pdf(file_path, document_folder)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            return self._extract_from_docx(file_path, document_folder)
        elif file_path.suffix.lower() in ['.pptx', '.ppt']:
            return self._extract_from_pptx(file_path, document_folder)
        elif file_path.suffix.lower() in ['.xlsx', '.xls']:
            return self._extract_from_excel(file_path, document_folder)
        elif file_path.suffix.lower() in ['.odt', '.ods', '.odp']:
            return self._extract_from_odf(file_path, document_folder)
        elif file_path.suffix.lower() in ['.html', '.htm', '.xml']:
            return self._extract_from_html_xml(file_path, document_folder)
        elif file_path.suffix.lower() in ['.zip', '.rar', '.7z']:
            return self._extract_from_archive(file_path, document_folder)
        else:
            return []

    def _create_document_folder(self, document_name: str) -> Path:
        """Create a unique folder for the document's images"""
        safe_name = "".join(c for c in document_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
        # Convert to lowercase and replace spaces and hyphens with underscores
        safe_name = safe_name.lower().replace(' ', '_').replace('-', '_')
        unique_id = str(uuid.uuid4())[:8]
        folder_name = f"{safe_name}_{unique_id}"
        document_folder = self.images_dir / folder_name
        document_folder.mkdir(parents=True, exist_ok=True)
        return document_folder

    def _extract_from_pdf(self, file_path: Path, output_folder: Path) -> List[ImageInfo]:
        """Extract images from PDF files with positioning information"""
        images = []
        try:
            pdf_document = fitz.open(file_path)

            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                image_list = page.get_images()

                # Extract text for this page to help with positioning context
                page_text = page.get_text()
                page_dict = page.get_text("dict")

                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    pix = fitz.Pixmap(pdf_document, xref)

                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        filename = f"page_{page_num + 1}_img_{img_index + 1}.png"
                        image_path = output_folder / filename
                        pix.save(str(image_path))

                        # Get image dimensions
                        with Image.open(image_path) as img_obj:
                            width, height = img_obj.size

                        # Get image position from page
                        image_rects = []
                        for block in page_dict.get("blocks", []):
                            if "image" in block:
                                image_rects.append(block["bbox"])

                        # Use the image index to get approximate position
                        position_x, position_y = None, None
                        if img_index < len(image_rects):
                            bbox = image_rects[img_index]
                            position_x = bbox[0]  # Left coordinate
                            position_y = bbox[1]  # Top coordinate

                        # Get surrounding text context for better positioning
                        content_context = self._get_text_context_around_image(page_text, position_y) if position_y else None

                        images.append(ImageInfo(
                            filename=filename,
                            url=f"{self.base_url}/images/{output_folder.name}/{filename}",
                            width=width,
                            height=height,
                            page_number=page_num + 1,
                            position_x=position_x,
                            position_y=position_y,
                            content_context=content_context
                        ))

                    pix = None

            pdf_document.close()
        except Exception as e:
            print(f"Error extracting images from PDF: {e}")

        return images

    def _get_text_context_around_image(self, page_text: str, image_y_position: float) -> str:
        """Extract text context around where an image appears on the page"""
        if not page_text or image_y_position is None:
            return None

        # Split text into lines and find context around the image position
        lines = page_text.split('\n')
        # For simplicity, return first few lines as context
        # In a more sophisticated implementation, we'd use coordinate data
        context_lines = lines[:3] if len(lines) >= 3 else lines
        return ' '.join(context_lines).strip()[:200]  # Limit context length

    def _extract_from_docx(self, file_path: Path, output_folder: Path) -> List[ImageInfo]:
        """Extract images from DOCX files with positioning information"""
        images = []
        try:
            doc = Document(file_path)

            # Create a mapping of image relationships to actual images
            image_rels = {}
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_rels[rel.rId] = rel.target_part.blob

            # Track processed images to avoid duplicates
            processed_images = set()
            image_count = 0

            # Walk through document elements to find images and their positions
            for paragraph_idx, paragraph in enumerate(doc.paragraphs):
                # Check for images in this paragraph using a simpler approach
                for run in paragraph.runs:
                    # Check if this run contains images by looking for drawing elements
                    if hasattr(run._element, 'xpath'):
                        try:
                            # Look for drawing elements (simpler xpath without namespaces)
                            drawings = run._element.xpath('.//w:drawing')
                            if drawings:
                                # Look for embedded image references
                                for drawing in drawings:
                                    # Get all attributes that might contain image references
                                    embed_ids = self._extract_embed_ids_from_drawing(drawing)

                                    for embed_id in embed_ids:
                                        if embed_id and embed_id in image_rels:
                                            # Create a unique identifier for this image data
                                            image_data = image_rels[embed_id]
                                            image_hash = hash(image_data)

                                            # Skip if we've already processed this exact image
                                            if image_hash in processed_images:
                                                print(f"DEBUG: Skipping duplicate image with embed_id {embed_id}")
                                                continue

                                            processed_images.add(image_hash)
                                            image_count += 1

                                            # Determine file extension from image data
                                            extension = self._get_image_extension(image_data)
                                            temp_filename = f"image_{image_count}.{extension}"
                                            temp_image_path = output_folder / temp_filename

                                            with open(temp_image_path, 'wb') as f:
                                                f.write(image_data)

                                            # Convert to PNG and cleanup original
                                            final_image_path = self._convert_to_png_and_cleanup(temp_image_path)
                                            final_filename = final_image_path.name

                                            # Get image dimensions
                                            try:
                                                with Image.open(final_image_path) as img_obj:
                                                    width, height = img_obj.size
                                            except:
                                                width, height = None, None

                                            # Create context from surrounding paragraphs
                                            content_context = self._get_docx_context(doc.paragraphs, paragraph_idx)

                                            # Calculate position in content (character-based estimation)
                                            position_in_content = self._estimate_content_position(doc.paragraphs, paragraph_idx)

                                            images.append(ImageInfo(
                                                filename=final_filename,
                                                url=f"{self.base_url}/images/{output_folder.name}/{final_filename}",
                                                width=width,
                                                height=height,
                                                position_in_content=position_in_content,
                                                content_context=content_context
                                            ))

                                            print(f"DEBUG: Found unique image {final_filename} at paragraph {paragraph_idx}, context: {content_context[:50] if content_context else 'None'}...")
                        except Exception as e:
                            print(f"DEBUG: Error processing run in paragraph {paragraph_idx}: {e}")
                            continue

            # If no images found through paragraph analysis, fall back to relationship method
            if not images:
                print("DEBUG: No images found in paragraph analysis, using relationship fallback")
                processed_rels = set()
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref and rel.rId not in processed_rels:
                        try:
                            processed_rels.add(rel.rId)
                            image_data = rel.target_part.blob
                            temp_filename = f"image_{len(images) + 1}.{rel.target_ref.split('.')[-1]}"
                            temp_image_path = output_folder / temp_filename

                            with open(temp_image_path, 'wb') as f:
                                f.write(image_data)

                            # Convert to PNG and cleanup original
                            final_image_path = self._convert_to_png_and_cleanup(temp_image_path)
                            final_filename = final_image_path.name

                            # Get image dimensions
                            try:
                                with Image.open(final_image_path) as img_obj:
                                    width, height = img_obj.size
                            except:
                                width, height = None, None

                            images.append(ImageInfo(
                                filename=final_filename,
                                url=f"{self.base_url}/images/{output_folder.name}/{final_filename}",
                                width=width,
                                height=height
                            ))
                        except Exception as e:
                            print(f"Error extracting image from DOCX: {e}")

        except Exception as e:
            print(f"Error processing DOCX file: {e}")

        return images

    def _extract_embed_ids_from_drawing(self, drawing_element) -> List[str]:
        """Extract embed IDs from a drawing element using string parsing as fallback"""
        embed_ids = []
        try:
            # Convert element to string and search for embed references
            element_str = str(drawing_element.xml) if hasattr(drawing_element, 'xml') else str(drawing_element)

            # Look for r:embed attributes in the XML
            import re
            embed_pattern = r'r:embed="([^"]+)"'
            matches = re.findall(embed_pattern, element_str)
            embed_ids.extend(matches)

            # Also look for embed attributes without namespace prefix
            embed_pattern2 = r'embed="([^"]+)"'
            matches2 = re.findall(embed_pattern2, element_str)
            embed_ids.extend(matches2)

        except Exception as e:
            print(f"DEBUG: Error extracting embed IDs: {e}")

        return embed_ids

    def _get_image_extension(self, image_data: bytes) -> str:
        """Determine image file extension from binary data"""
        if image_data.startswith(b'\x89PNG'):
            return 'png'
        elif image_data.startswith(b'\xff\xd8\xff'):
            return 'jpeg'
        elif image_data.startswith(b'GIF'):
            return 'gif'
        elif image_data.startswith(b'BM'):
            return 'bmp'
        else:
            return 'png'  # default fallback

    def _get_docx_context(self, paragraphs, current_idx: int) -> str:
        """Extract text context around the current paragraph position"""
        context_range = 2  # Look 2 paragraphs before and after
        start_idx = max(0, current_idx - context_range)
        end_idx = min(len(paragraphs), current_idx + context_range + 1)

        context_parts = []
        for i in range(start_idx, end_idx):
            if i < len(paragraphs):
                text = paragraphs[i].text.strip()
                if text and len(text) > 3:  # Only include meaningful text
                    context_parts.append(text)

        context = ' '.join(context_parts)
        return context[:200] if context else None  # Limit context length

    def _estimate_content_position(self, paragraphs, current_idx: int) -> int:
        """Estimate character position in the overall document content"""
        position = 0
        for i in range(current_idx):
            if i < len(paragraphs):
                position += len(paragraphs[i].text) + 1  # +1 for newline
        return position

    def _extract_from_pptx(self, file_path: Path, output_folder: Path) -> List[ImageInfo]:
        """Extract images from PPTX files"""
        images = []
        try:
            # PPTX files are ZIP archives
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                for file_info in zip_ref.filelist:
                    if file_info.filename.startswith('ppt/media/'):
                        filename = Path(file_info.filename).name
                        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff')):
                            image_data = zip_ref.read(file_info.filename)
                            temp_image_path = output_folder / filename

                            with open(temp_image_path, 'wb') as f:
                                f.write(image_data)

                            # Convert to PNG and cleanup original
                            final_image_path = self._convert_to_png_and_cleanup(temp_image_path)
                            final_filename = final_image_path.name

                            # Get image dimensions
                            try:
                                with Image.open(final_image_path) as img_obj:
                                    width, height = img_obj.size
                            except:
                                width, height = None, None

                            images.append(ImageInfo(
                                filename=final_filename,
                                url=f"{self.base_url}/images/{output_folder.name}/{final_filename}",
                                width=width,
                                height=height
                            ))
        except Exception as e:
            print(f"Error extracting images from PPTX: {e}")

        return images

    def _extract_from_excel(self, file_path: Path, output_folder: Path) -> List[ImageInfo]:
        """Extract images from Excel files"""
        images = []
        try:
            # Excel files are ZIP archives
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                for file_info in zip_ref.filelist:
                    if file_info.filename.startswith('xl/media/'):
                        filename = Path(file_info.filename).name
                        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff')):
                            image_data = zip_ref.read(file_info.filename)
                            temp_image_path = output_folder / filename

                            with open(temp_image_path, 'wb') as f:
                                f.write(image_data)

                            # Convert to PNG and cleanup original
                            final_image_path = self._convert_to_png_and_cleanup(temp_image_path)
                            final_filename = final_image_path.name

                            # Get image dimensions
                            try:
                                with Image.open(final_image_path) as img_obj:
                                    width, height = img_obj.size
                            except:
                                width, height = None, None

                            images.append(ImageInfo(
                                filename=final_filename,
                                url=f"{self.base_url}/images/{output_folder.name}/{final_filename}",
                                width=width,
                                height=height
                            ))
        except Exception as e:
            print(f"Error extracting images from Excel: {e}")

        return images

    def _extract_from_odf(self, file_path: Path, output_folder: Path) -> List[ImageInfo]:
        """Extract images from ODF files (OpenDocument Format)"""
        images = []
        try:
            # ODF files are ZIP archives
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                for file_info in zip_ref.filelist:
                    if file_info.filename.startswith('Pictures/'):
                        filename = Path(file_info.filename).name
                        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff')):
                            image_data = zip_ref.read(file_info.filename)
                            temp_image_path = output_folder / filename

                            with open(temp_image_path, 'wb') as f:
                                f.write(image_data)

                            # Convert to PNG and cleanup original
                            final_image_path = self._convert_to_png_and_cleanup(temp_image_path)
                            final_filename = final_image_path.name

                            # Get image dimensions
                            try:
                                with Image.open(final_image_path) as img_obj:
                                    width, height = img_obj.size
                            except:
                                width, height = None, None

                            images.append(ImageInfo(
                                filename=final_filename,
                                url=f"{self.base_url}/images/{output_folder.name}/{final_filename}",
                                width=width,
                                height=height
                            ))
        except Exception as e:
            print(f"Error extracting images from ODF: {e}")

        return images

    def _extract_from_html_xml(self, file_path: Path, output_folder: Path) -> List[ImageInfo]:
        """Extract images referenced in HTML/XML files"""
        images = []
        try:
            # For HTML/XML, we would need to parse and download referenced images
            # This is a basic implementation that looks for embedded base64 images
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # Look for base64 embedded images
            import re
            import base64

            pattern = r'data:image/([^;]+);base64,([^"]+)'
            matches = re.findall(pattern, content)

            for i, (image_type, base64_data) in enumerate(matches):
                try:
                    image_data = base64.b64decode(base64_data)
                    temp_filename = f"embedded_image_{i + 1}.{image_type}"
                    temp_image_path = output_folder / temp_filename

                    with open(temp_image_path, 'wb') as f:
                        f.write(image_data)

                    # Convert to PNG and cleanup original
                    final_image_path = self._convert_to_png_and_cleanup(temp_image_path)
                    final_filename = final_image_path.name

                    # Get image dimensions
                    try:
                        with Image.open(final_image_path) as img_obj:
                            width, height = img_obj.size
                    except:
                        width, height = None, None

                    images.append(ImageInfo(
                        filename=final_filename,
                        url=f"{self.base_url}/images/{output_folder.name}/{final_filename}",
                        width=width,
                        height=height
                    ))
                except Exception as e:
                    print(f"Error extracting embedded image: {e}")

        except Exception as e:
            print(f"Error processing HTML/XML file: {e}")

        return images

    def _extract_from_archive(self, file_path: Path, output_folder: Path) -> List[ImageInfo]:
        """Extract images from archive files"""
        images = []
        try:
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                for file_info in zip_ref.filelist:
                    filename = Path(file_info.filename).name
                    if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.svg')):
                        try:
                            image_data = zip_ref.read(file_info.filename)
                            safe_filename = "".join(c for c in filename if c.isalnum() or c in (' ', '-', '_', '.')).rstrip()
                            temp_image_path = output_folder / safe_filename

                            with open(temp_image_path, 'wb') as f:
                                f.write(image_data)

                            # Convert to PNG and cleanup original (skip SVG files)
                            if not filename.lower().endswith('.svg'):
                                final_image_path = self._convert_to_png_and_cleanup(temp_image_path)
                                final_filename = final_image_path.name

                                # Get image dimensions
                                try:
                                    with Image.open(final_image_path) as img_obj:
                                        width, height = img_obj.size
                                except:
                                    width, height = None, None
                            else:
                                # Keep SVG files as-is (no conversion)
                                final_filename = safe_filename
                                width, height = None, None

                            images.append(ImageInfo(
                                filename=final_filename,
                                url=f"{self.base_url}/images/{output_folder.name}/{final_filename}",
                                width=width,
                                height=height
                            ))
                        except Exception as e:
                            print(f"Error extracting image from archive: {e}")
        except Exception as e:
            print(f"Error processing archive file: {e}")

        return images

    def cleanup_old_images(self, days_old: int) -> Dict[str, Any]:
        """
        Delete image folders that are older than the specified number of days.

        Args:
            days_old: Number of days after which image folders should be deleted

        Returns:
            Dictionary with cleanup statistics
        """
        if not self.images_dir.exists():
            return {
                "status": "skipped",
                "reason": "Images directory does not exist",
                "deleted_folders": 0,
                "freed_space_mb": 0
            }

        cutoff_time = time.time() - (days_old * 24 * 60 * 60)  # Convert days to seconds
        deleted_folders = 0
        freed_space_bytes = 0
        deleted_folder_names = []

        try:
            for folder_path in self.images_dir.iterdir():
                if folder_path.is_dir():
                    # Get folder creation time
                    folder_stat = folder_path.stat()
                    folder_creation_time = folder_stat.st_ctime

                    if folder_creation_time < cutoff_time:
                        # Calculate folder size before deletion
                        folder_size = self._get_folder_size(folder_path)

                        try:
                            # Delete the folder and all its contents
                            shutil.rmtree(folder_path)
                            deleted_folders += 1
                            freed_space_bytes += folder_size
                            deleted_folder_names.append(folder_path.name)
                            print(f"Deleted old image folder: {folder_path.name}")
                        except Exception as e:
                            print(f"Error deleting folder {folder_path.name}: {e}")

            return {
                "status": "completed",
                "deleted_folders": deleted_folders,
                "freed_space_mb": round(freed_space_bytes / (1024 * 1024), 2),
                "deleted_folder_names": deleted_folder_names,
                "cleanup_time": datetime.datetime.now().isoformat()
            }

        except Exception as e:
            return {
                "status": "error",
                "error": str(e),
                "deleted_folders": deleted_folders,
                "freed_space_mb": round(freed_space_bytes / (1024 * 1024), 2)
            }

    def _get_folder_size(self, folder_path: Path) -> int:
        """Calculate the total size of a folder in bytes"""
        total_size = 0
        try:
            for file_path in folder_path.rglob('*'):
                if file_path.is_file():
                    total_size += file_path.stat().st_size
        except Exception as e:
            print(f"Error calculating size for {folder_path}: {e}")
        return total_size

    def _convert_to_png_and_cleanup(self, image_path: Path) -> Path:
        """Convert any image format to PNG and delete the original file"""
        try:
            # If it's already a PNG, return as-is
            if image_path.suffix.lower() == '.png':
                return image_path

            # Create new PNG filename
            png_path = image_path.with_suffix('.png')

            # Open and convert the image to PNG
            with Image.open(image_path) as img:
                # Convert to RGB if necessary (for formats like CMYK)
                if img.mode in ('RGBA', 'LA', 'P'):
                    # Keep transparency for formats that support it
                    img.save(png_path, 'PNG', optimize=True)
                elif img.mode in ('CMYK', 'YCbCr'):
                    # Convert CMYK and other modes to RGB first
                    rgb_img = img.convert('RGB')
                    rgb_img.save(png_path, 'PNG', optimize=True)
                else:
                    # For RGB and other modes, save directly
                    img.save(png_path, 'PNG', optimize=True)

            # Delete the original file
            if image_path != png_path and image_path.exists():
                image_path.unlink()
                print(f"Converted {image_path.name} to PNG and deleted original")

            return png_path

        except Exception as e:
            print(f"Error converting {image_path} to PNG: {e}")
            # If conversion fails, return the original path
            return image_path
